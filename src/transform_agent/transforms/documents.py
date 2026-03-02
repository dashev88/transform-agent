"""
Document transforms: PDF → text/Markdown, Excel ↔ JSON/CSV, DOCX → text/Markdown

Heavy imports (pymupdf, openpyxl, python-docx) are lazy-loaded to keep
cold-start fast for agents that only use text transforms.
"""

from __future__ import annotations

import io

import orjson
import polars as pl


# ---------------------------------------------------------------------------
# PDF → *
# ---------------------------------------------------------------------------

async def pdf_to_plain_text(data: bytes, options: dict | None = None) -> bytes:
    import pymupdf  # lazy import

    doc = pymupdf.open(stream=data, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages).encode()


async def pdf_to_markdown(data: bytes, options: dict | None = None) -> bytes:
    import pymupdf  # lazy import

    doc = pymupdf.open(stream=data, filetype="pdf")
    parts: list[str] = []
    for i, page in enumerate(doc):
        text = page.get_text()
        parts.append(f"## Page {i + 1}\n\n{text}")
    doc.close()
    return "\n\n---\n\n".join(parts).encode()


# ---------------------------------------------------------------------------
# Excel → *
# ---------------------------------------------------------------------------

async def excel_to_json(data: bytes, options: dict | None = None) -> bytes:
    sheet = (options or {}).get("sheet_name", None)
    # polars can read Excel directly
    df = pl.read_excel(io.BytesIO(data), sheet_name=sheet)
    rows = df.to_dicts()
    return orjson.dumps(rows, option=orjson.OPT_NON_STR_KEYS)


async def excel_to_csv(data: bytes, options: dict | None = None) -> bytes:
    sheet = (options or {}).get("sheet_name", None)
    df = pl.read_excel(io.BytesIO(data), sheet_name=sheet)
    buf = io.BytesIO()
    df.write_csv(buf)
    return buf.getvalue()


async def excel_to_html(data: bytes, options: dict | None = None) -> bytes:
    sheet = (options or {}).get("sheet_name", None)
    df = pl.read_excel(io.BytesIO(data), sheet_name=sheet)
    # Build simple HTML table
    rows = df.to_dicts()
    if not rows:
        return b"<table></table>"
    cols = list(rows[0].keys())
    lines = ["<table>", "<thead><tr>"]
    for c in cols:
        lines.append(f"<th>{c}</th>")
    lines.append("</tr></thead><tbody>")
    for row in rows:
        lines.append("<tr>")
        for c in cols:
            lines.append(f"<td>{row.get(c, '')}</td>")
        lines.append("</tr>")
    lines.append("</tbody></table>")
    return "\n".join(lines).encode()


# ---------------------------------------------------------------------------
# DOCX → *
# ---------------------------------------------------------------------------

async def docx_to_plain_text(data: bytes, options: dict | None = None) -> bytes:
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(paragraphs).encode()


async def docx_to_markdown(data: bytes, options: dict | None = None) -> bytes:
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        text = p.text.strip()
        if not text:
            continue
        if "heading 1" in style:
            parts.append(f"# {text}")
        elif "heading 2" in style:
            parts.append(f"## {text}")
        elif "heading 3" in style:
            parts.append(f"### {text}")
        elif "list" in style:
            parts.append(f"- {text}")
        else:
            parts.append(text)
    return "\n\n".join(parts).encode()


# ---------------------------------------------------------------------------
# JSON → Excel  (tabular.py handles CSV→Excel; we add JSON→Excel here)
# ---------------------------------------------------------------------------

async def json_to_excel(data: bytes, options: dict | None = None) -> bytes:
    obj = orjson.loads(data)
    if isinstance(obj, dict):
        obj = [obj]
    if not isinstance(obj, list):
        obj = [{"value": obj}]
    df = pl.DataFrame(obj)
    buf = io.BytesIO()
    df.write_excel(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# JSON / CSV → HTML table
# ---------------------------------------------------------------------------

async def json_to_html(data: bytes, options: dict | None = None) -> bytes:
    obj = orjson.loads(data)
    if isinstance(obj, dict):
        obj = [obj]
    if not isinstance(obj, list) or not obj:
        return b"<pre>" + orjson.dumps(obj, option=orjson.OPT_INDENT_2) + b"</pre>"
    cols = list(obj[0].keys()) if isinstance(obj[0], dict) else ["value"]
    lines = ["<table>", "<thead><tr>"]
    for c in cols:
        lines.append(f"<th>{c}</th>")
    lines.append("</tr></thead><tbody>")
    for row in obj:
        lines.append("<tr>")
        if isinstance(row, dict):
            for c in cols:
                lines.append(f"<td>{row.get(c, '')}</td>")
        else:
            lines.append(f"<td>{row}</td>")
        lines.append("</tr>")
    lines.append("</tbody></table>")
    return "\n".join(lines).encode()


async def json_to_markdown_table(data: bytes, options: dict | None = None) -> bytes:
    obj = orjson.loads(data)
    if isinstance(obj, dict):
        obj = [obj]
    if not isinstance(obj, list) or not obj:
        return orjson.dumps(obj, option=orjson.OPT_INDENT_2)
    if not isinstance(obj[0], dict):
        return "\n".join(str(r) for r in obj).encode()
    cols = list(obj[0].keys())
    lines = ["| " + " | ".join(cols) + " |"]
    lines.append("| " + " | ".join("---" for _ in cols) + " |")
    for row in obj:
        vals = [str(row.get(c, "")) for c in cols]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines).encode()


async def csv_to_html(data: bytes, options: dict | None = None) -> bytes:
    sep = (options or {}).get("delimiter", ",")
    df = pl.read_csv(io.BytesIO(data), separator=sep, infer_schema_length=1000)
    rows = df.to_dicts()
    if not rows:
        return b"<table></table>"
    cols = list(rows[0].keys())
    lines = ["<table>", "<thead><tr>"]
    for c in cols:
        lines.append(f"<th>{c}</th>")
    lines.append("</tr></thead><tbody>")
    for row in rows:
        lines.append("<tr>")
        for c in cols:
            lines.append(f"<td>{row.get(c, '')}</td>")
        lines.append("</tr>")
    lines.append("</tbody></table>")
    return "\n".join(lines).encode()


async def csv_to_markdown(data: bytes, options: dict | None = None) -> bytes:
    sep = (options or {}).get("delimiter", ",")
    df = pl.read_csv(io.BytesIO(data), separator=sep, infer_schema_length=1000)
    rows = df.to_dicts()
    if not rows:
        return b""
    cols = list(rows[0].keys())
    lines = ["| " + " | ".join(cols) + " |"]
    lines.append("| " + " | ".join("---" for _ in cols) + " |")
    for row in rows:
        vals = [str(row.get(c, "")) for c in cols]
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines).encode()
